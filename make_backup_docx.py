"""
make_backup_docx.py
===================
원본 6개 docx를 받아 본문에 박힌 샘플 이름을 {NAME} 변수로
일괄 치환한 백업 docx 6개를 생성한다.

[입력] ./docs/ 안의 원본 6개 docx
       ./sample_names.txt 안의 샘플 이름 목록 (한 줄에 한 이름)
[출력] ./backup_docx/ 안에 _백업.docx 6개

치환 규칙:
  [샘플이름]님 → {NAME}님
  [샘플이름]   → {NAME}

원본의 서식(굵게/단락/제목)은 그대로 유지.
"""

from docx import Document
from pathlib import Path
import re

INPUT_DIR = Path("./docs")
OUTPUT_DIR = Path("./backup_docx")
OUTPUT_DIR.mkdir(exist_ok=True)

# 샘플 이름 외부 파일에서 로드
SAMPLE_NAMES_FILE = Path("./sample_names.txt")
if SAMPLE_NAMES_FILE.exists():
    SAMPLE_NAMES = [
        line.strip() for line in SAMPLE_NAMES_FILE.read_text(encoding='utf-8').splitlines()
        if line.strip() and not line.strip().startswith('#')
    ]
else:
    SAMPLE_NAMES = []
    print("⚠️  sample_names.txt 없음 — 치환 작업 건너뜀")


def replace_in_runs(paragraph):
    """단락 안의 모든 run에서 치환 수행."""
    full_text = paragraph.text
    if not any(w in full_text for w in SAMPLE_NAMES):
        return False
    
    new_text = full_text
    for sample in SAMPLE_NAMES:
        new_text = re.sub(re.escape(sample) + r'님', '{NAME}님', new_text)
        new_text = re.sub(re.escape(sample), '{NAME}', new_text)
    
    if new_text == full_text:
        return False
    
    runs = paragraph.runs
    if not runs:
        return False
    
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''
    return True


def process_docx(in_path, out_path):
    """docx 1개를 읽어 치환 후 저장."""
    doc = Document(in_path)
    changes = 0
    
    for para in doc.paragraphs:
        if replace_in_runs(para):
            changes += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_runs(para):
                        changes += 1
    
    doc.save(out_path)
    return changes


def main():
    files = [
        '카드해설_1차_메이저정방향.docx',
        '카드해설_2차_메이저역방향.docx',
        '카드해설_3차_Wands.docx',
        '카드해설_4차_Cups.docx',
        'Swords_5차_본문_448개.docx',
        '카드해설_6차_Pentacles.docx',
    ]
    
    if not SAMPLE_NAMES:
        print("샘플 이름 목록이 비어 있어 처리할 작업 없음.")
        return
    
    for fname in files:
        in_path = INPUT_DIR / fname
        stem = in_path.stem
        out_name = f"{stem}_백업.docx"
        out_path = OUTPUT_DIR / out_name
        
        if not in_path.exists():
            print(f"⚠️  파일 없음: {in_path}")
            continue
        
        changes = process_docx(in_path, out_path)
        size_kb = out_path.stat().st_size / 1024
        print(f"✅ {fname:42s} → {out_name:46s}  ({size_kb:6.1f}KB, 치환 {changes}건)")
    
    print(f"\n📁 백업 docx 저장 위치: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
